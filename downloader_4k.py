"""
YouTube 4K/8K Bulk Downloader from Word File
===========================================
Extracts YouTube links from a Word document (.docx) and downloads
them in the highest available quality by merging separate video
and audio streams using ffmpeg.

Requirements:
    - python-docx
    - pytubefix
    - ffmpeg (must be installed separately)

Usage:
    1. Place your Word file (videos.docx) in the same folder
    2. Update FFMPEG_PATH below to match your system
    3. Run: python downloader_4k.py
    4. Type 'y' when prompted to start downloads

Author: Created for personal use
Date: 2026-07-03
"""

from docx import Document
from pytubefix import YouTube
from pytubefix.cli import on_progress
import os
import re
import subprocess


# ==================== CONFIGURATION ====================
# UPDATE THESE VALUES FOR YOUR SYSTEM

WORD_FILE = 'videos.docx'           # Your Word file with YouTube links
OUTPUT_FOLDER = 'downloads'         # Where videos will be saved

# FFMPEG PATH - UPDATE THIS FOR YOUR SYSTEM
# Windows example:  r"D:\youtube_downloader\ffmpeg-8.1.2-essentials_build\ffmpeg-8.1.2-essentials_build\bin\ffmpeg.exe"
# Mac example:      "/usr/local/bin/ffmpeg"
# Linux example:    "/usr/bin/ffmpeg"
FFMPEG_PATH = 'ffmpeg'

# =======================================================


def extract_youtube_links_from_docx(file_path):
    """
    Extract all YouTube links from a Word document.
    Searches both hyperlinks and plain text URLs.
    
    Args:
        file_path (str): Path to .docx file
    
    Returns:
        list: Unique YouTube URLs found in the document
    """
    doc = Document(file_path)
    youtube_links = []
    
    # Regex pattern to match YouTube URLs and extract 11-char video IDs
    youtube_pattern = re.compile(
        r'(https?://)?(www\.)?(youtube|youtu)\.(com|be)/'
        r'(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})'
    )
    
    # Search in all paragraphs
    for paragraph in doc.paragraphs:
        # Check plain text for URLs
        text = paragraph.text
        matches = youtube_pattern.findall(text)
        for match in matches:
            video_id = match[5]
            url = f"https://www.youtube.com/watch?v={video_id}"
            if url not in youtube_links:
                youtube_links.append(url)
        
        # Check hyperlinks in runs (clickable links)
        for run in paragraph.runs:
            if hasattr(run, 'hyperlink') and run.hyperlink:
                if run.hyperlink.address:
                    url = run.hyperlink.address
                    if 'youtube.com' in url or 'youtu.be' in url:
                        if url not in youtube_links:
                            youtube_links.append(url)
    
    # Search in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = youtube_pattern.findall(text)
                for match in matches:
                    video_id = match[5]
                    url = f"https://www.youtube.com/watch?v={video_id}"
                    if url not in youtube_links:
                        youtube_links.append(url)
    
    return youtube_links


def check_ffmpeg():
    """
    Verify ffmpeg is installed and accessible.
    
    Returns:
        bool: True if ffmpeg found, False otherwise
    """
    try:
        result = subprocess.run(
            [FFMPEG_PATH, '-version'],
            capture_output=True,
            text=True,
            timeout=10
        )
        if result.returncode == 0:
            version = result.stdout.split('\n')[0]
            print(f"✅ ffmpeg found: {version[:50]}...")
            return True
    except FileNotFoundError:
        pass
    except Exception as e:
        pass
    
    print("❌ ffmpeg NOT found!")
    print(f"   Path checked: {FFMPEG_PATH}")
    print("   Please install ffmpeg and update FFMPEG_PATH in the code.")
    print("   Download: https://www.gyan.dev/ffmpeg/builds/")
    return False


def get_best_streams(yt):
    """
    Get highest quality video and audio streams separately.
    YouTube stores high-res video and audio in separate adaptive streams.
    
    Args:
        yt (YouTube): pytubefix YouTube object
    
    Returns:
        tuple: (video_stream, audio_stream)
    """
    # Highest resolution video-only stream (no audio)
    video_stream = yt.streams.filter(
        adaptive=True, only_video=True
    ).order_by('resolution').desc().first()
    
    # Highest bitrate audio-only stream
    audio_stream = yt.streams.filter(
        adaptive=True, only_audio=True
    ).order_by('abr').desc().first()
    
    return video_stream, audio_stream


def download_and_merge(yt, output_path, video_stream, audio_stream):
    """
    Download video and audio separately, then merge with ffmpeg.
    Uses -c copy for zero quality loss (no re-encoding).
    
    Args:
        yt (YouTube): pytubefix YouTube object
        output_path (str): Folder to save final video
        video_stream: Video-only stream object
        audio_stream: Audio-only stream object
    
    Returns:
        str: Path to final merged file, or None if failed
    """
    # Clean filename - remove invalid characters
    title = re.sub(r'[<>:"/\\|?*]', '', yt.title)[:100]
    
    # Temporary file paths
    temp_video = os.path.join(output_path, f"temp_video_{title}.mp4")
    temp_audio = os.path.join(output_path, f"temp_audio_{title}.mp4")
    final_output = os.path.join(output_path, f"{title}.mp4")
    
    # Download video stream
    print(f"   📹 Downloading video ({video_stream.resolution})...")
    video_stream.download(
        output_path=output_path,
        filename=f"temp_video_{title}.mp4"
    )
    
    # Download audio stream
    print(f"   🔊 Downloading audio ({audio_stream.abr})...")
    audio_stream.download(
        output_path=output_path,
        filename=f"temp_audio_{title}.mp4"
    )
    
    # Merge using ffmpeg (copy codec = no quality loss, very fast)
    print(f"   🔗 Merging video + audio with ffmpeg...")
    
    cmd = [
        FFMPEG_PATH,
        '-i', temp_video,           # Input 1: video file
        '-i', temp_audio,           # Input 2: audio file
        '-c', 'copy',               # Copy streams without re-encoding
        '-map', '0:v:0',            # Map video from input 0
        '-map', '1:a:0',            # Map audio from input 1
        '-shortest',                # End when shortest stream ends
        '-y',                       # Overwrite output if exists
        final_output
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    # Clean up temporary files regardless of success/failure
    if os.path.exists(temp_video):
        os.remove(temp_video)
    if os.path.exists(temp_audio):
        os.remove(temp_audio)
    
    if result.returncode != 0:
        print(f"   ❌ ffmpeg merge failed!")
        print(f"   Error: {result.stderr[:300]}")
        return None
    
    return final_output


def download_video_highest_quality(url, output_path='downloads'):
    """
    Download a single YouTube video in highest available quality.
    
    Args:
        url (str): YouTube video URL
        output_path (str): Folder to save video
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        print(f"\n{'='*60}")
        print(f"📥 URL: {url}")
        
        # Initialize YouTube object with progress bar callback
        yt = YouTube(url, on_progress_callback=on_progress)
        
        print(f"🎬 Title: {yt.title}")
        print(f"⏱️  Length: {yt.length // 60}m {yt.length % 60}s")
        print(f"👤 Channel: {yt.author}")
        
        # Create output directory if it doesn't exist
        os.makedirs(output_path, exist_ok=True)
        
        # Get best separate streams
        video_stream, audio_stream = get_best_streams(yt)
        
        # Fallback to progressive stream if adaptive not available
        if not video_stream or not audio_stream:
            print("   ⚠️  Adaptive streams not available, using progressive...")
            stream = yt.streams.get_highest_resolution()
            file_path = stream.download(output_path=output_path)
            print(f"\n✅ SAVED: {os.path.basename(file_path)}")
            return True
        
        # Display stream information
        print(f"🎯 Video: {video_stream.resolution} ({video_stream.mime_type})")
        print(f"🎯 Audio: {audio_stream.abr} ({audio_stream.mime_type})")
        
        # Calculate total download size
        total_size = (video_stream.filesize or 0) + (audio_stream.filesize or 0)
        print(f"💾 Total Size: ~{total_size / (1024*1024):.1f} MB")
        
        # Download and merge
        final_path = download_and_merge(yt, output_path, video_stream, audio_stream)
        
        if final_path:
            print(f"\n✅ SAVED: {os.path.basename(final_path)}")
            return True
        else:
            return False
            
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        return False


def main():
    """Main execution function."""
    print("=" * 70)
    print("🎬 YOUTUBE 4K/8K HIGHEST QUALITY DOWNLOADER")
    print("=" * 70)
    
    # Step 1: Verify ffmpeg is available
    if not check_ffmpeg():
        return
    
    # Step 2: Verify Word file exists
    if not os.path.exists(WORD_FILE):
        print(f"\n❌ ERROR: File '{WORD_FILE}' not found!")
        print(f"   Current folder: {os.getcwd()}")
        print(f"\n   Files in this folder:")
        for f in os.listdir('.'):
            print(f"      - {f}")
        return
    
    # Step 3: Extract links from Word document
    print(f"\n📄 Reading: {WORD_FILE}")
    links = extract_youtube_links_from_docx(WORD_FILE)
    
    if not links:
        print("❌ No YouTube links found! Check your Word file.")
        return
    
    print(f"\n✅ Found {len(links)} link(s):")
    for i, link in enumerate(links, 1):
        print(f"   {i}. {link}")
    
    # Step 4: Analyze videos (show estimated sizes)
    print(f"\n📊 Analyzing videos...")
    total_estimated = 0
    for i, link in enumerate(links, 1):
        try:
            yt = YouTube(link)
            video_stream, audio_stream = get_best_streams(yt)
            if video_stream and audio_stream:
                size_mb = ((video_stream.filesize or 0) + (audio_stream.filesize or 0)) / (1024*1024)
                total_estimated += size_mb
                print(f"   [{i}] {yt.title[:45]}... | {video_stream.resolution} | ~{size_mb:.1f} MB")
            else:
                stream = yt.streams.get_highest_resolution()
                size_mb = (stream.filesize or 0) / (1024*1024)
                total_estimated += size_mb
                print(f"   [{i}] {yt.title[:45]}... | {stream.resolution} | ~{size_mb:.1f} MB")
        except Exception as e:
            print(f"   [{i}] Could not analyze: {str(e)[:50]}")
    
    print(f"\n💾 Estimated total download: ~{total_estimated:.1f} MB ({total_estimated/1024:.1f} GB)")
    
    # Step 5: Confirm before downloading
    confirm = input(f"\n⚡ Download all {len(links)} videos in HIGHEST QUALITY? (y/n): ")
    if confirm.lower() != 'y':
        print("❌ Cancelled.")
        return
    
    # Step 6: Start downloads
    print("\n" + "=" * 70)
    
    success = 0
    failed = 0
    for i, link in enumerate(links, 1):
        print(f"\n{'='*70}")
        print(f"🔄 [{i}/{len(links)}]")
        if download_video_highest_quality(link, OUTPUT_FOLDER):
            success += 1
        else:
            failed += 1
    
    # Step 7: Final summary
    print("\n" + "=" * 70)
    print("📊 DOWNLOAD SUMMARY")
    print("=" * 70)
    print(f"✅ Successful: {success}")
    print(f"❌ Failed: {failed}")
    print(f"📁 Saved to: {os.path.abspath(OUTPUT_FOLDER)}")
    print("=" * 70)


if __name__ == "__main__":
    main()
