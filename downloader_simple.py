"""
YouTube Bulk Downloader (Simple Version - No ffmpeg Required)
=============================================================
Extracts YouTube links from a Word document and downloads them
in the best available progressive quality (up to 720p).

No ffmpeg needed - uses built-in pytubefix progressive streams.

Requirements:
    - python-docx
    - pytubefix

Usage:
    1. Place videos.docx in the same folder
    2. Run: python downloader_simple.py
"""

from docx import Document
from pytubefix import YouTube
from pytubefix.cli import on_progress
import os
import re


# ==================== CONFIGURATION ====================
WORD_FILE = 'videos.docx'
OUTPUT_FOLDER = 'downloads'
# =======================================================


def extract_youtube_links_from_docx(file_path):
    """Extract YouTube links from Word document."""
    doc = Document(file_path)
    youtube_links = []
    
    youtube_pattern = re.compile(
        r'(https?://)?(www\.)?(youtube|youtu)\.(com|be)/'
        r'(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})'
    )
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = youtube_pattern.findall(text)
        for match in matches:
            video_id = match[5]
            url = f"https://www.youtube.com/watch?v={video_id}"
            if url not in youtube_links:
                youtube_links.append(url)
        
        for run in paragraph.runs:
            if hasattr(run, 'hyperlink') and run.hyperlink:
                if run.hyperlink.address:
                    url = run.hyperlink.address
                    if 'youtube.com' in url or 'youtu.be' in url:
                        if url not in youtube_links:
                            youtube_links.append(url)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = youtube_pattern.findall(text)
                for match in matches:
                    video_id = match[5]
                    url = f"https://www.youtube.com/watch?v={video_id}"
                    if url not in youtube_links:
                        youtube_links.append(url)
    
    return youtube_links


def download_video(url, output_path='downloads'):
    """Download video in highest progressive quality."""
    try:
        print(f"\n{'='*50}")
        print(f"📥 URL: {url}")
        
        yt = YouTube(url, on_progress_callback=on_progress)
        
        print(f"🎬 Title: {yt.title}")
        print(f"⏱️  Length: {yt.length // 60}m {yt.length % 60}s")
        print(f"👤 Channel: {yt.author}")
        
        os.makedirs(output_path, exist_ok=True)
        
        # Get highest progressive stream (video + audio combined)
        stream = yt.streams.get_highest_resolution()
        
        print(f"🎯 Quality: {stream.resolution}")
        print(f"💾 Size: ~{stream.filesize / (1024*1024):.1f} MB")
        print(f"⬇️  Downloading...")
        
        file_path = stream.download(output_path=output_path)
        filename = os.path.basename(file_path)
        
        print(f"✅ SAVED: {filename}")
        return True
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        return False


def main():
    print("=" * 60)
    print("🎬 YOUTUBE DOWNLOADER (Simple - No ffmpeg)")
    print("=" * 60)
    
    if not os.path.exists(WORD_FILE):
        print(f"\n❌ ERROR: '{WORD_FILE}' not found!")
        print(f"   Current folder: {os.getcwd()}")
        return
    
    print(f"\n📄 Reading: {WORD_FILE}")
    links = extract_youtube_links_from_docx(WORD_FILE)
    
    if not links:
        print("❌ No YouTube links found!")
        return
    
    print(f"\n✅ Found {len(links)} link(s):")
    for i, link in enumerate(links, 1):
        print(f"   {i}. {link}")
    
    confirm = input(f"\n⚡ Download all {len(links)} videos? (y/n): ")
    if confirm.lower() != 'y':
        print("❌ Cancelled.")
        return
    
    print("\n" + "=" * 60)
    
    success = 0
    failed = 0
    for i, link in enumerate(links, 1):
        print(f"\n🔄 [{i}/{len(links)}]")
        if download_video(link, OUTPUT_FOLDER):
            success += 1
        else:
            failed += 1
    
    print("\n" + "=" * 60)
    print(f"✅ Done: {success} | ❌ Failed: {failed}")
    print(f"📁 Saved to: {os.path.abspath(OUTPUT_FOLDER)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
